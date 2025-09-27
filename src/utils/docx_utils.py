from pathlib import Path

import docx


def extract_docx_text(file_path: Path | str) -> str:
    path = Path(file_path)
    doc = docx.Document(str(path))
    return "\n".join([p.text for p in doc.paragraphs])


def extract_docx_metadata(file_path: Path | str) -> dict:
    path = Path(file_path)
    doc = docx.Document(str(path))
    core = doc.core_properties
    return {
        "title": core.title,
        "author": core.author,
        "created": str(core.created),
    }
